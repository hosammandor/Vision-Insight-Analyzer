import streamlit as st
import google.generativeai as genai
import fitz  # PyMuPDF
from PIL import Image
import io
import requests
import pandas as pd
from docx import Document

# --- إعدادات الصفحة ---
st.set_page_config(page_title="Vision Insight Power", page_icon="🚀", layout="wide")

# --- CSS التصميم الاحترافي ---
st.markdown("""
    <style>
    .stApp { background: radial-gradient(circle at top right, #1e1b4b, #0f172a); color: #f8fafc; }
    div[data-testid="stVerticalBlock"] > div:has(div.status-box) {
        background: rgba(255, 255, 255, 0.03); backdrop-filter: blur(10px);
        border-radius: 20px; border: 1px solid rgba(255, 255, 255, 0.1); padding: 25px;
    }
    .stButton>button {
        background: linear-gradient(135deg, #8b5cf6 0%, #d946ef 100%);
        color: white; border: none; border-radius: 12px; font-weight: 700; height: 3em;
    }
    .download-btn { margin-top: 10px; }
    .status-box { background: rgba(15, 23, 42, 0.6); padding: 20px; border-radius: 15px; border-right: 4px solid #a855f7; }
    </style>
    """, unsafe_allow_html=True)

# --- وظائف التحويل للملفات ---
def create_word_doc(text):
    doc = Document()
    doc.add_heading('نتائج تحليل Vision Insight', 0)
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_excel_from_text(text):
    # محاولة بسيطة لتحويل الجداول النصية (Markdown) إلى DataFrame
    try:
        from io import StringIO
        # البحث عن جداول مارك داون في النص
        if "|" in text:
            # تنظيف النص لاستخراج الجدول فقط (تبسيط)
            lines = [line for line in text.split('\n') if "|" in line]
            table_str = '\n'.join(lines)
            df = pd.read_csv(StringIO(table_str.replace(' ', '')), sep="|").dropna(axis=1, how='all')
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Sheet1')
            return output.getvalue()
    except:
        return None
    return None

# --- Sidebar ---
with st.sidebar:
    st.title("🔮 Power Panel")
    api_key = st.text_input("Gemini API Key:", type="password")
    if api_key:
        genai.configure(api_key=api_key)
        model_name = st.selectbox("Model:", ["gemini-1.5-pro", "gemini-1.5-flash"])

# --- واجهة المستخدم ---
st.markdown("<h1 style='text-align: center;'>Vision <span style='color: #a855f7;'>Insight</span> Power</h1>", unsafe_allow_html=True)

if api_key:
    model = genai.GenerativeModel(model_name)
    col1, col2 = st.columns([1, 1.2], gap="large")
    
    with col1:
        st.markdown("### 📥 المستند أو الرابط")
        input_mode = st.radio("المصدر:", ["ملف", "رابط"])
        content_imgs = []
        
        if input_mode == "ملف":
            up = st.file_uploader("", type=["pdf", "png", "jpg", "jpeg"])
            if up:
                if up.type == "application/pdf":
                    doc = fitz.open(stream=up.read(), filetype="pdf")
                    for p in doc:
                        pix = p.get_pixmap(matrix=fitz.Matrix(2, 2))
                        content_imgs.append(Image.open(io.BytesIO(pix.tobytes("png"))))
                else:
                    img = Image.open(up)
                    content_imgs.append(img)
                    st.image(img, use_container_width=True)
        else:
            url = st.text_input("رابط الصورة:")
            if url:
                r = requests.get(url)
                img = Image.open(io.BytesIO(r.content))
                content_imgs.append(img)
                st.image(img, use_container_width=True)

    with col2:
        st.markdown("### 🤖 التحليل والتحويل")
        q = st.text_area("ماذا نفعل بالبيانات؟", placeholder="مثال: استخرج الجدول المالي، أو اكتب تقريراً كاملاً...")
        
        if st.button("تنفيذ 🚀") and content_imgs:
            with st.spinner("جاري المعالجة..."):
                res = model.generate_content([q] + content_imgs)
                st.session_state['last_res'] = res.text
                st.markdown(f'<div class="status-box">{res.text}</div>', unsafe_allow_html=True)

        if 'last_res' in st.session_state:
            st.markdown("### 📥 تحميل النتائج:")
            c1, c2 = st.columns(2)
            
            # زر الوورد
            word_data = create_word_doc(st.session_state['last_res'])
            c1.download_button("تحميل ملف Word 📄", word_data, "analysis.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            # زر الإكسيل (يظهر إذا وجد جدول)
            excel_data = create_excel_from_text(st.session_state['last_res'])
            if excel_data:
                c2.download_button("تحميل ملف Excel 📊", excel_data, "data.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            else:
                c2.info("لم يتم العثور على جدول لتحويله لـ Excel")
else:
    st.warning("الرجاء إدخال الـ API Key")
